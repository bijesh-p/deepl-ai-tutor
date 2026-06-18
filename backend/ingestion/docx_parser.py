from __future__ import annotations

import uuid
from pathlib import Path

import docx

from backend.ingestion.models import Document, Section, SourceType

_DEFAULT_MAX_SECTIONS = 16
_HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}
_HEADING_LEVEL = {"Heading 1": 1, "Heading 2": 2, "Heading 3": 3, "Heading 4": 4}


def parse_docx(file_path: str, max_sections: int = _DEFAULT_MAX_SECTIONS) -> Document:
    """Parse a Word document into the unified Document model.

    Sections are derived from heading paragraphs. Body paragraphs between
    headings are accumulated into the preceding section's text. If no headings
    exist, the entire document becomes one section. At most *max_sections*
    sections are produced.
    """
    path = Path(file_path)
    doc_id = str(uuid.uuid4())

    d = docx.Document(str(path))

    raw_title = (d.core_properties.title or "").strip()
    title = raw_title if raw_title else path.stem

    sections = _extract_sections(d, max_sections)
    if not sections:
        # No headings and no body text — produce a minimal fallback section
        all_text = "\n\n".join(p.text.strip() for p in d.paragraphs if p.text.strip())
        if all_text:
            sections = [Section(section_id=str(uuid.uuid4()), title=title, body=all_text, level=1)]

    return Document(
        doc_id=doc_id,
        title=title,
        source_filename=path.name,
        source_type=SourceType.DOCX,
        sections=sections,
        total_pages=len(sections),
    )


def _extract_sections(d: docx.Document, max_sections: int) -> list[Section]:
    sections: list[Section] = []
    current_title: str | None = None
    current_level: int = 1
    body_parts: list[str] = []

    def _flush() -> None:
        if current_title is None:
            return
        body = "\n\n".join(body_parts)
        if body:
            sections.append(
                Section(
                    section_id=str(uuid.uuid4()),
                    title=current_title,
                    body=body,
                    level=current_level,
                )
            )

    for para in d.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        if style_name in _HEADING_STYLES:
            _flush()
            if len(sections) >= max_sections:
                break
            current_title = text or style_name
            current_level = _HEADING_LEVEL.get(style_name, 1)
            body_parts = []
        else:
            if text:
                body_parts.append(text)

    _flush()

    # If no headings were found, produce a single section from all paragraphs
    if not sections:
        all_text = "\n\n".join(p.text.strip() for p in d.paragraphs if p.text.strip())
        if all_text:
            sections = [Section(section_id=str(uuid.uuid4()), title="Content", body=all_text, level=1)]

    return sections[:max_sections]
